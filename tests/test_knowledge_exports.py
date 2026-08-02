import io
import tempfile
import unittest
from contextlib import contextmanager
from pathlib import Path

from docx import Document

import app as geo_app
from services.auth import create_user
from services.knowledge_exports import build_knowledge_docx


@contextmanager
def isolated_export_app():
    original = {
        "D": geo_app.D,
        "F_CLIENTS": geo_app.F_CLIENTS,
        "F_USERS": geo_app.F_USERS,
        "F_GROUPS": geo_app.F_GROUPS,
        "AUTH_DISABLED": geo_app.app.config.get("AUTH_DISABLED"),
    }
    with tempfile.TemporaryDirectory() as tmp:
        geo_app.D = tmp
        geo_app.F_CLIENTS = str(Path(tmp) / "clients.json")
        geo_app.F_USERS = str(Path(tmp) / "users.json")
        geo_app.F_GROUPS = str(Path(tmp) / "probe_groups.json")
        geo_app.app.config["AUTH_DISABLED"] = False
        try:
            yield Path(tmp)
        finally:
            for key, value in original.items():
                if key == "AUTH_DISABLED":
                    if value is None:
                        geo_app.app.config.pop(key, None)
                    else:
                        geo_app.app.config[key] = value
                else:
                    setattr(geo_app, key, value)


class KnowledgeDocxExportTests(unittest.TestCase):
    def test_builds_readable_docx_from_markdown_sections(self):
        content = "# 客户总资料\n\n## 产品与服务\n\n- 成人学历提升\n- 全流程提醒\n"
        output = build_knowledge_docx("客户资料知识库", content, "测试客户")

        document = Document(io.BytesIO(output.getvalue()))
        paragraphs = [paragraph.text for paragraph in document.paragraphs]
        self.assertIn("客户资料知识库", paragraphs)
        self.assertIn("产品与服务", paragraphs)
        self.assertIn("成人学历提升", paragraphs)
        self.assertEqual(document.paragraphs[0].style.name, "Title")

    def test_owner_can_download_four_allowed_knowledge_types_only(self):
        with isolated_export_app() as root:
            create_user(geo_app.F_USERS, "owner", "secret-pass", role="operator")
            geo_app.save(geo_app.F_CLIENTS, [{"id": "client-a", "owner_username": "owner", "name": "测试客户", "industry": "教育"}])
            geo_app.save(geo_app.F_GROUPS, {"client-a": [{"id": "group-a", "name": "问题组", "questions": ["问题一"]}]})
            knowledge_dir = root / "knowledge_base" / "client-a"
            knowledge_dir.mkdir(parents=True)
            (knowledge_dir / "customer_master.md").write_text("# 客户总资料\n\n## 产品与服务\n\n客户资料。\n", encoding="utf-8")
            (knowledge_dir / "competitor_master.md").write_text("# 竞品总资料\n\n## 竞品甲\n\n竞品资料。\n", encoding="utf-8")

            client = geo_app.app.test_client()
            client.post("/api/auth/login", json={"username": "owner", "password": "secret-pass"})
            for kind in ("customer", "competitors", "routes", "scenes"):
                response = client.get(f"/api/knowledge/export/client-a/{kind}")
                self.assertEqual(response.status_code, 200)
                self.assertEqual(response.headers["Content-Type"].split(";")[0], "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                self.assertTrue(response.data.startswith(b"PK"))
                self.assertIn("attachment", response.headers["Content-Disposition"])
            self.assertEqual(client.get("/api/knowledge/export/client-a/quality").status_code, 404)
            self.assertEqual(client.get("/api/knowledge/export/client-a/system-prompts").status_code, 404)
