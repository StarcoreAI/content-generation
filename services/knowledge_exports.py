"""DOCX exports for the operator-facing knowledge libraries."""
import io
import re

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn


DOCX_MIMETYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


def _set_font(style, name="Microsoft YaHei", size=11, color=None, bold=None):
    font = style.font
    font.name = name
    font.size = Pt(size)
    if color:
        font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        font.bold = bold
    style.element.rPr.rFonts.set(qn("w:ascii"), name)
    style.element.rPr.rFonts.set(qn("w:hAnsi"), name)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), name)


def _configure_document(document):
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = document.styles["Normal"]
    _set_font(normal)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    title = document.styles["Title"]
    _set_font(title, size=18, color="2E74B5", bold=True)
    title.paragraph_format.space_after = Pt(10)

    for level, size, before, after, color in (
        (1, 16, 18, 10, "2E74B5"),
        (2, 13, 14, 7, "2E74B5"),
        (3, 12, 10, 5, "1F4D78"),
    ):
        style = document.styles[f"Heading {level}"]
        _set_font(style, size=size, color=color, bold=True)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    for style_name in ("List Bullet", "List Number"):
        style = document.styles[style_name]
        _set_font(style)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    if "Knowledge metadata" not in document.styles:
        metadata = document.styles.add_style("Knowledge metadata", WD_STYLE_TYPE.PARAGRAPH)
        _set_font(metadata, size=9, color="666666")
        metadata.paragraph_format.space_after = Pt(12)


def _append_markdown(document, content):
    lines = str(content or "").replace("\r\n", "\n").split("\n")
    for raw_line in lines:
        line = raw_line.strip()
        if not line:
            continue
        heading = re.match(r"^(#{1,6})\s+(.+?)\s*$", line)
        if heading:
            level = min(max(len(heading.group(1)), 2), 4) - 1
            document.add_paragraph(heading.group(2), style=f"Heading {level}")
            continue
        bullet = re.match(r"^[-*]\s+(.+)$", line)
        if bullet:
            document.add_paragraph(bullet.group(1), style="List Bullet")
            continue
        numbered = re.match(r"^\d+[.)]\s+(.+)$", line)
        if numbered:
            document.add_paragraph(numbered.group(1), style="List Number")
            continue
        document.add_paragraph(line)


def build_knowledge_docx(title, content, client_name=""):
    """Return an in-memory, readable DOCX for a single knowledge library."""
    document = Document()
    _configure_document(document)
    document.add_paragraph(str(title or "知识库"), style="Title")
    if client_name:
        document.add_paragraph(f"客户：{client_name}", style="Knowledge metadata")
    if str(content or "").strip():
        _append_markdown(document, content)
    else:
        document.add_paragraph("暂无已沉淀资料。")
    output = io.BytesIO()
    document.save(output)
    output.seek(0)
    return output


def routes_markdown(industry, groups):
    chunks = ["# 行业写法库"]
    if industry:
        chunks.extend(["", f"行业：{industry}"])
    for group_name in ("介绍型", "对比型"):
        chunks.extend(["", f"## {group_name}"])
        routes = (groups or {}).get(group_name) or []
        if not routes:
            chunks.extend(["", "暂无已沉淀路线。"])
            continue
        for route in routes:
            chunks.extend(["", f"### {route.get('name') or '未命名路线'}"])
            for field, label in (("reader_task", "读者任务"), ("signature", "路线特征"), ("risk_notes", "适用边界")):
                value = str(route.get(field) or "").strip()
                if value:
                    chunks.append(f"{label}：{value}")
            for step in route.get("steps") or []:
                purpose = str(step.get("purpose") or "").strip()
                action = str(step.get("output_action") or "").strip()
                if purpose or action:
                    chunks.append(f"- {purpose}：{action}".rstrip("："))
    return "\n".join(chunks).strip() + "\n"


def scenes_markdown(rows):
    chunks = ["# 场景词库"]
    grouped = {}
    for row in rows or []:
        group_name = str(row.get("group_name") or row.get("group_id") or "未分组").strip()
        grouped.setdefault(group_name, []).append(row)
    for group_name, group_rows in grouped.items():
        chunks.extend(["", f"## {group_name}"])
        for row in group_rows:
            query = str(row.get("query") or "未命名问题").strip()
            terms = [str(item).strip() for item in row.get("scene_terms") or [] if str(item).strip()]
            chunks.extend(["", f"### {query}", "", "、".join(terms) if terms else "暂无场景词。"])
    return "\n".join(chunks).strip() + "\n"
